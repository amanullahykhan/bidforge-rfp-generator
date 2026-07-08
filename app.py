import os
import json
import sqlite3
import threading
import requests
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
import customtkinter as ctk

# Try to import python-docx for document parsing/exporting
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ==========================================
# Database Manager (Complex History DB)
# ==========================================
class HistoryDB:
    def __init__(self, db_path="bidforge_history.db"):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS generations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                    provider TEXT,
                    tone TEXT,
                    rfp_prompt TEXT,
                    response TEXT
                )
            ''')
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS context_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    generation_id INTEGER,
                    filename TEXT,
                    FOREIGN KEY(generation_id) REFERENCES generations(id)
                )
            ''')
            conn.commit()

    def save_generation(self, provider, tone, prompt, response, context_filenames):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO generations (provider, tone, rfp_prompt, response)
                VALUES (?, ?, ?, ?)
            ''', (provider, tone, prompt, response))
            
            generation_id = cursor.lastrowid
            
            for filename in context_filenames:
                cursor.execute('''
                    INSERT INTO context_files (generation_id, filename)
                    VALUES (?, ?)
                ''', (generation_id, filename))
            conn.commit()

# ==========================================
# LLM Client (Unified API Router)
# ==========================================
class LLMClient:
    def __init__(self, config):
        self.config = config

    def generate(self, provider, rfp_text, context_text, tone):
        api_key = self.config.get("api_keys", {}).get(provider, "")
        base_url = self.config.get("base_urls", {}).get(provider, "")
        
        system_prompt = (
            f"You are a Principal Technical Proposal Writer. Your task is to draft an RFP response. "
            f"Use the provided company context to accurately represent the company's capabilities. "
            f"Adopt a {tone} tone. Do not invent false metrics; strictly rely on the context provided."
        )
        
        user_prompt = f"--- COMPANY CONTEXT ---\n{context_text}\n\n--- RFP REQUIREMENTS ---\n{rfp_text}"

        try:
            if provider == "OpenAI":
                return self._call_openai(api_key, system_prompt, user_prompt)
            elif provider == "Anthropic":
                return self._call_anthropic(api_key, system_prompt, user_prompt)
            elif provider == "Gemini":
                return self._call_gemini(api_key, system_prompt, user_prompt)
            elif provider == "DeepSeek":
                return self._call_deepseek(api_key, system_prompt, user_prompt)
            elif provider == "Local Ollama":
                return self._call_ollama(base_url, system_prompt, user_prompt)
            else:
                return "Unknown provider selected."
        except Exception as e:
            return f"Error connecting to {provider}: {str(e)}"

    def _call_openai(self, api_key, system, user):
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": "gpt-4o",
            "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}]
        }
        res = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        res.raise_for_status()
        return res.json()["choices"][0]["message"]["content"]

    def _call_anthropic(self, api_key, system, user):
        headers = {"x-api-key": api_key, "anthropic-version": "2023-06-01", "Content-Type": "application/json"}
        payload = {
            "model": "claude-3-opus-20240229",
            "system": system,
            "messages": [{"role": "user", "content": user}],
            "max_tokens": 4096
        }
        res = requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=payload)
        res.raise_for_status()
        return res.json()["content"][0]["text"]

    def _call_gemini(self, api_key, system, user):
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro:generateContent?key={api_key}"
        payload = {
            "system_instruction": {"parts": [{"text": system}]},
            "contents": [{"parts": [{"text": user}]}]
        }
        res = requests.post(url, json=payload)
        res.raise_for_status()
        return res.json()["candidates"][0]["content"]["parts"][0]["text"]

    def _call_deepseek(self, api_key, system, user):
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": "deepseek-chat",
            "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}]
        }
        res = requests.post("https://api.deepseek.com/chat/completions", headers=headers, json=payload)
        res.raise_for_status()
        return res.json()["choices"][0]["message"]["content"]

    def _call_ollama(self, base_url, system, user):
        url = f"{base_url.rstrip('/')}/api/generate"
        payload = {
            "model": "qwen2.5-coder:7b",
            "prompt": f"{system}\n\n{user}",
            "stream": False
        }
        res = requests.post(url, json=payload)
        res.raise_for_status()
        return res.json()["response"]


# ==========================================
# Main GUI Application
# ==========================================
class BidForgeApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # --- App Config & Theme ---
        self.title("BidForge AI - Enterprise RFP Generator")
        self.geometry("1100x700")
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        self.config_file = "config.json"
        self.app_config = self.load_config()
        self.history_db = HistoryDB()
        self.llm_client = LLMClient(self.app_config)
        self.knowledge_bank = {} # filename: text

        self._build_ui()
        self._load_ui_state()

    def load_config(self):
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, "r") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {"api_keys": {}, "base_urls": {"Local Ollama": "http://localhost:11434"}}

    def save_config(self):
        provider = self.provider_var.get()
        api_key = self.api_key_entry.get().strip()
        base_url = self.base_url_entry.get().strip()
        
        if "api_keys" not in self.app_config:
            self.app_config["api_keys"] = {}
        if "base_urls" not in self.app_config:
            self.app_config["base_urls"] = {}
            
        self.app_config["api_keys"][provider] = api_key
        self.app_config["base_urls"][provider] = base_url

        with open(self.config_file, "w") as f:
            json.dump(self.app_config, f, indent=4)
            
        self.llm_client.config = self.app_config
        self.status_label.configure(text=f"Config saved for {provider}", text_color="#00FF00")
        self.after(3000, lambda: self.status_label.configure(text="Ready", text_color="gray"))

    def _build_ui(self):
        # Grid Layout Strategy
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        # --- Left Sidebar (Config) ---
        self.sidebar_frame = ctk.CTkFrame(self, width=250, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(6, weight=1)

        ctk.CTkLabel(self.sidebar_frame, text="BidForge AI", font=ctk.CTkFont(size=24, weight="bold")).grid(row=0, column=0, padx=20, pady=(20, 10))
        ctk.CTkLabel(self.sidebar_frame, text="API Configuration", font=ctk.CTkFont(size=14)).grid(row=1, column=0, padx=20, pady=(10, 5), sticky="w")

        self.provider_var = ctk.StringVar(value="OpenAI")
        self.provider_dropdown = ctk.CTkOptionMenu(self.sidebar_frame, variable=self.provider_var, 
                                                   values=["OpenAI", "Anthropic", "Gemini", "DeepSeek", "Local Ollama"],
                                                   command=self._on_provider_change)
        self.provider_dropdown.grid(row=2, column=0, padx=20, pady=10, sticky="ew")

        self.api_key_entry = ctk.CTkEntry(self.sidebar_frame, placeholder_text="API Key", show="*")
        self.api_key_entry.grid(row=3, column=0, padx=20, pady=10, sticky="ew")

        self.base_url_entry = ctk.CTkEntry(self.sidebar_frame, placeholder_text="Local Base URL")
        self.base_url_entry.grid(row=4, column=0, padx=20, pady=10, sticky="ew")
        self.base_url_entry.grid_remove() # Hide initially

        self.save_btn = ctk.CTkButton(self.sidebar_frame, text="Save Config", command=self.save_config)
        self.save_btn.grid(row=5, column=0, padx=20, pady=20, sticky="ew")

        self.status_label = ctk.CTkLabel(self.sidebar_frame, text="Ready", text_color="gray")
        self.status_label.grid(row=7, column=0, padx=20, pady=20, sticky="s")

        # --- Main Window (Tabview) ---
        self.tabview = ctk.CTkTabview(self, corner_radius=10)
        self.tabview.grid(row=0, column=1, padx=20, pady=20, sticky="nsew")
        
        self.tab1 = self.tabview.add("Knowledge Bank")
        self.tab2 = self.tabview.add("RFP Generator")
        self.tab3 = self.tabview.add("Output Review")

        self._build_tab1()
        self._build_tab2()
        self._build_tab3()

    def _build_tab1(self):
        self.tab1.grid_columnconfigure(0, weight=1)
        self.tab1.grid_rowconfigure(1, weight=1)

        # "Drag and Drop" styled upload zone
        self.upload_frame = ctk.CTkFrame(self.tab1, corner_radius=15, border_width=2, border_color="#333333")
        self.upload_frame.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        self.upload_frame.bind("<Button-1>", lambda e: self._add_documents())
        
        upload_lbl = ctk.CTkLabel(self.upload_frame, text="Drop Documents Here or Click to Browse\n(.txt, .docx)", font=ctk.CTkFont(size=16))
        upload_lbl.pack(pady=40, padx=40)
        upload_lbl.bind("<Button-1>", lambda e: self._add_documents())

        # Indexed Files List
        list_frame = ctk.CTkFrame(self.tab1)
        list_frame.grid(row=1, column=0, padx=20, pady=(0, 20), sticky="nsew")
        list_frame.grid_columnconfigure(0, weight=1)
        list_frame.grid_rowconfigure(1, weight=1)

        ctk.CTkLabel(list_frame, text="Currently Indexed Files:", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, padx=10, pady=10, sticky="w")
        
        self.file_listbox = ctk.CTkScrollableFrame(list_frame)
        self.file_listbox.grid(row=1, column=0, padx=10, pady=(0, 10), sticky="nsew")

        self.clear_index_btn = ctk.CTkButton(list_frame, text="Clear Index", fg_color="#C62828", hover_color="#B71C1C", command=self._clear_index)
        self.clear_index_btn.grid(row=2, column=0, padx=10, pady=10, sticky="e")

    def _build_tab2(self):
        self.tab2.grid_columnconfigure(0, weight=1)
        self.tab2.grid_rowconfigure(1, weight=1)

        header_frame = ctk.CTkFrame(self.tab2, fg_color="transparent")
        header_frame.grid(row=0, column=0, sticky="ew")
        header_frame.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(header_frame, text="Paste RFP Requirements / Questions:", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, sticky="w", pady=(0, 10))
        
        self.tone_var = ctk.StringVar(value="Executive")
        self.tone_dropdown = ctk.CTkOptionMenu(header_frame, variable=self.tone_var, values=["Executive", "Technical", "Conservative", "Persuasive"])
        self.tone_dropdown.grid(row=0, column=1, sticky="e", pady=(0, 10))

        self.rfp_input = ctk.CTkTextbox(self.tab2, wrap="word", font=ctk.CTkFont(size=14))
        self.rfp_input.grid(row=1, column=0, sticky="nsew", pady=(0, 20))

        self.generate_btn = ctk.CTkButton(self.tab2, text="Draft Response", height=50, font=ctk.CTkFont(size=16, weight="bold"), command=self._start_generation)
        self.generate_btn.grid(row=2, column=0, sticky="ew")
        
        self.progress_bar = ctk.CTkProgressBar(self.tab2, mode="indeterminate")
        self.progress_bar.grid(row=3, column=0, sticky="ew", pady=(10, 0))
        self.progress_bar.grid_remove() # Hide initially

    def _build_tab3(self):
        self.tab3.grid_columnconfigure(0, weight=1)
        self.tab3.grid_rowconfigure(0, weight=1)

        self.output_text = ctk.CTkTextbox(self.tab3, wrap="word", font=ctk.CTkFont(size=14))
        self.output_text.grid(row=0, column=0, sticky="nsew", pady=(0, 20))

        btn_frame = ctk.CTkFrame(self.tab3, fg_color="transparent")
        btn_frame.grid(row=1, column=0, sticky="ew")
        btn_frame.grid_columnconfigure((0, 1), weight=1)

        self.copy_btn = ctk.CTkButton(btn_frame, text="Copy to Clipboard", command=self._copy_to_clipboard)
        self.copy_btn.grid(row=0, column=0, padx=(0, 10), sticky="ew")

        self.export_btn = ctk.CTkButton(btn_frame, text="Export to Docx", command=self._export_to_docx)
        self.export_btn.grid(row=0, column=1, padx=(10, 0), sticky="ew")

    def _load_ui_state(self):
        self._on_provider_change(self.provider_var.get())

    def _on_provider_change(self, choice):
        self.api_key_entry.delete(0, tk.END)
        self.base_url_entry.delete(0, tk.END)

        if choice == "Local Ollama":
            self.api_key_entry.grid_remove()
            self.base_url_entry.grid()
            url = self.app_config.get("base_urls", {}).get(choice, "http://localhost:11434")
            self.base_url_entry.insert(0, url)
        else:
            self.base_url_entry.grid_remove()
            self.api_key_entry.grid()
            key = self.app_config.get("api_keys", {}).get(choice, "")
            self.api_key_entry.insert(0, key)

    # --- Knowledge Bank Logic ---
    def _add_documents(self):
        filetypes = [("Text/Word Files", "*.txt *.docx"), ("All Files", "*.*")]
        filepaths = filedialog.askopenfilenames(title="Select Knowledge Bank Documents", filetypes=filetypes)
        
        for path in filepaths:
            filename = os.path.basename(path)
            if filename in self.knowledge_bank:
                continue
            
            try:
                text = ""
                if path.endswith(".txt"):
                    with open(path, "r", encoding="utf-8") as f:
                        text = f.read()
                elif path.endswith(".docx") and DOCX_AVAILABLE:
                    doc = docx.Document(path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                elif path.endswith(".docx") and not DOCX_AVAILABLE:
                    messagebox.showwarning("Missing Dependency", f"Cannot parse {filename}. python-docx is not installed.")
                    continue
                
                self.knowledge_bank[filename] = text
                lbl = ctk.CTkLabel(self.file_listbox, text=f"📄 {filename}", anchor="w")
                lbl.pack(fill="x", padx=5, pady=2)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to read {filename}:\n{str(e)}")

    def _clear_index(self):
        self.knowledge_bank.clear()
        for widget in self.file_listbox.winfo_children():
            widget.destroy()

    # --- Generation Logic ---
    def _start_generation(self):
        rfp_prompt = self.rfp_input.get("1.0", tk.END).strip()
        if not rfp_prompt:
            messagebox.showwarning("Input Required", "Please enter RFP Requirements before generating.")
            return

        # UI Updates
        self.generate_btn.configure(state="disabled", text="Generating response, please wait...")
        self.progress_bar.grid()
        self.progress_bar.start()
        self.output_text.delete("1.0", tk.END)

        provider = self.provider_var.get()
        tone = self.tone_var.get()
        
        # Compile Context
        context = "\n\n---\n\n".join([f"Document: {fname}\n{text}" for fname, text in self.knowledge_bank.items()])
        if not context:
            context = "No company context provided."

        # Start Thread
        threading.Thread(target=self._generation_thread, args=(provider, rfp_prompt, context, tone), daemon=True).start()

    def _generation_thread(self, provider, rfp_prompt, context, tone):
        response = self.llm_client.generate(provider, rfp_prompt, context, tone)
        
        # Save to DB history
        context_files = list(self.knowledge_bank.keys())
        self.history_db.save_generation(provider, tone, rfp_prompt, response, context_files)

        # Update GUI safely
        self.after(0, lambda: self._finish_generation(response))

    def _finish_generation(self, response):
        self.progress_bar.stop()
        self.progress_bar.grid_remove()
        self.generate_btn.configure(state="normal", text="Draft Response")
        
        self.output_text.insert("1.0", response)
        self.tabview.set("Output Review")

    # --- Output Logic ---
    def _copy_to_clipboard(self):
        text = self.output_text.get("1.0", tk.END).strip()
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)
            self.copy_btn.configure(text="Copied!")
            self.after(2000, lambda: self.copy_btn.configure(text="Copy to Clipboard"))

    def _export_to_docx(self):
        if not DOCX_AVAILABLE:
            messagebox.showerror("Dependency Error", "python-docx is not installed. Run: pip install python-docx")
            return

        text = self.output_text.get("1.0", tk.END).strip()
        if not text:
            return

        filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], title="Export RFP Response")
        if filepath:
            try:
                doc = docx.Document()
                doc.add_heading("RFP Response Draft", 0)
                doc.add_paragraph(text)
                doc.save(filepath)
                messagebox.showinfo("Success", f"Document saved to {filepath}")
            except Exception as e:
                messagebox.showerror("Export Error", f"Failed to save document:\n{str(e)}")


if __name__ == "__main__":
    app = BidForgeApp()
    app.mainloop()
